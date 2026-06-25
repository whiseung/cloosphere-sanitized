"""create_docx — Word 문서 생성 툴."""

import logging
from io import BytesIO
from typing import Literal, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from extension_modules.tools.document._common import (
    SYSTEM_TOKEN_KEYS,
    TOKEN_PATTERN,
    format_tool_response,
    load_docx_template,
    save_to_files,
    system_token_values,
)
from langchain_core.tools import StructuredTool
from pydantic import BaseModel, Field, model_validator

log = logging.getLogger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_HEADER_FILL_HEX = "D9E2F3"  # 옅은 푸른빛 회색 — 비즈니스 보고서 톤


def _set_cell_bg(cell, color_hex: str) -> None:
    """Word 테이블 셀 배경색 (XML 직접 조작 — python-docx 정식 API 미제공)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


class DocxBlock(BaseModel):
    type: Literal["heading", "paragraph", "bullet", "table"]
    level: Optional[int] = Field(None, ge=1, le=3, description="heading 전용")
    text: Optional[str] = Field(None, description="heading/paragraph/bullet 본문")
    table: Optional[list[list[str]]] = Field(None, description="첫 행=헤더, table 전용")

    @model_validator(mode="after")
    def _validate_block(self):
        if self.type == "heading":
            if not self.text or self.level is None:
                raise ValueError("heading은 text와 level 필수")
        elif self.type in ("paragraph", "bullet"):
            if not self.text:
                raise ValueError(f"{self.type}은 text 필수")
        elif self.type == "table":
            if not self.table or len(self.table) < 1:
                raise ValueError("table은 최소 1행 필수")
        return self


class DocxContent(BaseModel):
    filename: str = Field(..., description="확장자 제외, 한글/영문 가능")
    title: str = Field(..., description="문서 첫 줄에 들어갈 제목")
    template_tokens: Optional[dict[str, str]] = Field(
        None,
        description=(
            "관리자 Word 템플릿 표지의 `{{key}}` 자리표시자를 런타임에 치환할 값. "
            "예: `{'title': '2026 Q1 매출 보고서'}`. `date`/`author` 등 시스템 토큰은 "
            "자동 채워지므로 LLM 이 컨텐츠 토큰만 신경 쓰면 된다. 키 누락 시 "
            "`{{key}}` 가 그대로 남는다. 템플릿 미사용 시 무시됨."
        ),
    )
    blocks: list[DocxBlock] = Field(..., min_length=1, max_length=500)


# ── Builder ──────────────────────────────────────────────────────────────────


_W_TXBX_CONTENT = qn("w:txbxContent")
_W_P = qn("w:p")


def _iter_textbox_paragraphs(part_element):
    """주어진 XML element (body 또는 header/footer) 안의 모든 textbox/shape
    paragraph yield.

    Word 의 표지 디자인은 흔히 텍스트 상자(도형)로 만들어짐 — 자유 배치를 위해서.
    그런데 ``doc.paragraphs`` / ``cell.paragraphs`` 같은 python-docx 표준 API 는
    flow content (body, table cell) 의 paragraph 만 반환하고 textbox/shape 안의
    paragraph 는 안 보여줌 (관측 사례: 사용자 표지의 `{{title}}` 가 DrawingML +
    VML textbox 안에 있어 _iter_doc_paragraphs 가 못 찾고 치환 실패).

    DrawingML (`<w:drawing>...<w:txbxContent>`) 와 VML (`<v:textbox>...<w:txbxContent>`)
    둘 다 결국 `<w:txbxContent>` 안에 `<w:p>` 가 들어가므로 XML 트리에서 직접 찾는다.
    """
    from docx.text.paragraph import Paragraph

    for txbx in part_element.iter(_W_TXBX_CONTENT):
        for p_elem in txbx.iter(_W_P):
            # parent=None — substitute 는 run.text/p.text 만 만져 part-level 참조 불필요
            yield Paragraph(p_elem, None)


def _iter_doc_paragraphs(doc):
    """문서 body + 테이블 셀 + section header/footer + 모든 textbox 안의 paragraph yield.

    토큰 검색/치환 시 누락 영역 없도록 — 표지 디자인은 textbox/shape 에 분산
    되는 경우가 많음 (`{{title}}` 가 도형 안인 케이스 관측).
    """
    # Body
    for p in doc.paragraphs:
        yield p
    # Body 테이블
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                # 중첩 테이블도 한 단계 deep
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for p in ncell.paragraphs:
                                yield p
    # Body 영역의 textbox/shape 안 paragraph
    yield from _iter_textbox_paragraphs(doc.element.body)
    # Section header/footer
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                for p in hf.paragraphs:
                    yield p
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p
                # header/footer 안 textbox 도 스캔
                yield from _iter_textbox_paragraphs(hf._element)
            except Exception:  # noqa: BLE001 — header/footer 미정의 시 skip
                continue


def _collect_docx_tokens(doc) -> set[str]:
    """Word 문서의 모든 paragraph 텍스트에서 발견된 `{{key}}` 키 set."""
    keys: set[str] = set()
    for p in _iter_doc_paragraphs(doc):
        if "{{" in p.text:
            keys.update(TOKEN_PATTERN.findall(p.text))
    return keys


def _discover_docx_template_token_keys() -> set[str]:
    """현재 활성 DOCX 템플릿을 로드해 발견된 모든 토큰 키 set 반환 (description 용)."""
    try:
        doc = load_docx_template()
        if doc is None:
            return set()
        return _collect_docx_tokens(doc)
    except Exception as e:  # noqa: BLE001
        log.warning("Failed to discover DOCX template tokens: %s", e)
        return set()


def _substitute_docx_tokens(doc, tokens: dict[str, str]) -> None:
    """문서의 모든 paragraph 에서 `{{key}}` → tokens[key] 치환.

    paragraph 단위로 처리하면 형식(폰트/색/굵기)이 깨지는 경우가 있어, 가능한 한
    run 단위로 처리. 단 토큰이 여러 run 에 걸쳐 split 된 케이스 (예: '{{tit' 한 run
    + 'le}}' 다른 run) 는 run 합치고 첫 run 에 결과 넣고 나머지 비움 — 흔한 케이스.

    누락 키는 원본 `{{key}}` 유지.
    """
    if not tokens:
        return

    def _replace(match):
        return tokens.get(match.group(1), match.group(0))

    for p in _iter_doc_paragraphs(doc):
        if "{{" not in p.text:
            continue
        # 1) Run-단위 시도 — 가장 흔한 케이스 (토큰이 단일 run 안에 있는 경우).
        # 형식 (폰트/색/굵기) 보존됨.
        for run in p.runs:
            if run.text and "{{" in run.text:
                new_text = TOKEN_PATTERN.sub(_replace, run.text)
                if new_text != run.text:
                    run.text = new_text
        # 2) Paragraph.text 에 토큰이 남아있으면 (Word 가 토큰을 여러 run 으로 split
        # 한 케이스) — paragraph 전체 치환. 첫 run 에 합친 결과 넣고 나머지 비움.
        # 형식 일부 소실 trade-off — 토큰이 split 안 됐다면 1단계에서 다 처리됨.
        if "{{" in p.text:
            full_new = TOKEN_PATTERN.sub(_replace, p.text)
            if full_new != p.text:
                if p.runs:
                    p.runs[0].text = full_new
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = full_new


def _build_docx(
    content: DocxContent,
    *,
    system_tokens: Optional[dict[str, str]] = None,
) -> BytesIO:
    """DocxContent → in-memory .docx BytesIO.

    템플릿 미설정 시: 'Title', 'Heading 1', 'List Bullet', 'Table Grid' 내장 스타일 + 헤더 셀 옅은 푸른빛 회색 (#D9E2F3).
    템플릿 설정 시: 마스터 docx 의 styles.xml inherit — 'Title' / 'Heading' 등이 템플릿 정의대로 적용되고 기존 콘텐츠(양식 헤더 등) 뒤에 새 본문이 append 된다.
        헤더 셀 배경 hard-coded 는 skip 해 마스터 표 스타일에 맡긴다.

    토큰 치환:
    - 템플릿의 `{{key}}` 자리표시자 → ``content.template_tokens`` + ``system_tokens``
      (LLM 값이 시스템 값보다 우선)
    - 표지 토큰 적용 시 ``content.title`` 의 추가 paragraph 는 skip (cover 가 이미 title
      역할). 토큰 미적용 시 기존 동작 (Title 스타일 paragraph append) 유지.
    """
    template = load_docx_template()
    template_used = template is not None
    doc = template if template_used else Document()

    cover_token_substituted = False
    if template_used:
        merged_tokens: dict[str, str] = {
            **(system_tokens or {}),
            "title": content.title,  # LLM title fallback — admin 이 {{title}} 박았으면 채워짐
            **(content.template_tokens or {}),
        }
        if merged_tokens:
            discovered = _collect_docx_tokens(doc)
            unfilled = discovered - merged_tokens.keys()
            if unfilled:
                log.info(
                    "create_docx template has unfilled tokens %s — left as-is",
                    sorted(unfilled),
                )
            _substitute_docx_tokens(doc, merged_tokens)
            # template_tokens 가 명시 전달된 경우 = admin 이 표지 디자인 의도 → 중복 title 차단
            if content.template_tokens or "title" in discovered:
                cover_token_substituted = True

    if not cover_token_substituted:
        doc.add_paragraph(content.title, style="Title")

    for block in content.blocks:
        if block.type == "heading":
            doc.add_heading(block.text, level=block.level)
        elif block.type == "paragraph":
            doc.add_paragraph(block.text)
        elif block.type == "bullet":
            doc.add_paragraph(block.text, style="List Bullet")
        elif block.type == "table":
            rows = block.table
            n_cols = len(rows[0]) if rows else 0
            # 'Table Grid' = 모든 셀 border 가 있는 내장 스타일
            table = doc.add_table(rows=len(rows), cols=n_cols, style="Table Grid")
            for r_idx, row in enumerate(rows):
                for c_idx, cell_value in enumerate(row):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = str(cell_value)
                    if r_idx == 0:
                        # 헤더 셀: bold + 배경색 (배경은 템플릿 미사용 시만)
                        if not template_used:
                            _set_cell_bg(cell, _HEADER_FILL_HEX)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


_TOOL_DESCRIPTION = (
    "Generate a Word document (.docx) file. "
    "보고서/문서를 .docx 파일로 생성한다. "
    "Use when the user asks for: 'Word document', 'Word doc', 'docx', 'report', "
    "'document file', '워드 문서', '보고서', '문서 파일', '계약서 초안' etc. "
    "Pass an ordered sequence of blocks (heading / paragraph / bullet / table)."
)


def _build_dynamic_description() -> str:
    """기본 description + 현재 DOCX 템플릿의 컨텐츠 토큰 키 안내."""
    base = _TOOL_DESCRIPTION
    discovered = _discover_docx_template_token_keys()
    content_keys = sorted(discovered - SYSTEM_TOKEN_KEYS)
    if content_keys:
        keys_csv = ", ".join(content_keys)
        base += (
            " IMPORTANT — the admin's Word template has these content placeholders "
            f"that you should fill via `template_tokens`: {keys_csv}. "
            "Date/author tokens are auto-filled by the system."
        )
    return base


def make_create_docx(user_id: str) -> StructuredTool:
    """user_id에 바인딩된 create_docx 툴을 생성."""

    def _create_docx(**kwargs) -> str:
        log.info(
            "create_docx called: user_id=%s, template_tokens=%r, block_count=%d",
            user_id,
            kwargs.get("template_tokens"),
            len(kwargs.get("blocks") or []),
        )
        content = DocxContent(**kwargs)
        system_tokens = system_token_values(user_id)
        buf = _build_docx(content, system_tokens=system_tokens)
        result = save_to_files(
            user_id=user_id,
            filename=f"{content.filename}.docx",
            buffer=buf,
            mime=DOCX_MIME,
        )
        return format_tool_response(result)

    return StructuredTool.from_function(
        func=_create_docx,
        name="create_docx",
        description=_build_dynamic_description(),
        args_schema=DocxContent,
    )
