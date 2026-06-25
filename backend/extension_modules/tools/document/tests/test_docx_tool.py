"""create_docx 툴 테스트."""

from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxOpen
from docx.oxml.ns import qn
from extension_modules.tools.document.docx_tool import (
    DOCX_MIME,
    DocxBlock,
    DocxContent,
    _build_docx,
    make_create_docx,
)
from pydantic import ValidationError


class TestDocxSchema:
    def test_heading_requires_text_and_level(self):
        with pytest.raises(ValidationError):
            DocxBlock(type="heading")
        with pytest.raises(ValidationError):
            DocxBlock(type="heading", text="t")  # level 없음
        b = DocxBlock(type="heading", text="t", level=1)
        assert b.level == 1

    def test_heading_level_range(self):
        with pytest.raises(ValidationError):
            DocxBlock(type="heading", text="t", level=0)
        with pytest.raises(ValidationError):
            DocxBlock(type="heading", text="t", level=4)

    def test_paragraph_requires_text(self):
        with pytest.raises(ValidationError):
            DocxBlock(type="paragraph")
        b = DocxBlock(type="paragraph", text="hello")
        assert b.text == "hello"

    def test_bullet_requires_text(self):
        with pytest.raises(ValidationError):
            DocxBlock(type="bullet")
        b = DocxBlock(type="bullet", text="item")
        assert b.text == "item"

    def test_table_requires_table(self):
        with pytest.raises(ValidationError):
            DocxBlock(type="table")
        b = DocxBlock(type="table", table=[["h1", "h2"], ["a", "b"]])
        assert len(b.table) == 2

    def test_empty_blocks_rejected(self):
        with pytest.raises(ValidationError):
            DocxContent(filename="r", title="T", blocks=[])

    def test_minimal_valid(self):
        c = DocxContent(
            filename="r",
            title="My Report",
            blocks=[DocxBlock(type="paragraph", text="Hello")],
        )
        assert len(c.blocks) == 1


class TestDocxBuilder:
    def test_round_trip_all_block_types(self):
        content = DocxContent(
            filename="r",
            title="Report",
            blocks=[
                DocxBlock(type="heading", text="Section 1", level=1),
                DocxBlock(type="paragraph", text="Body text."),
                DocxBlock(type="bullet", text="Item A"),
                DocxBlock(type="bullet", text="Item B"),
                DocxBlock(type="table", table=[["H1", "H2"], ["a", "b"]]),
            ],
        )
        buf = _build_docx(content)
        buf.seek(0)
        doc = DocxOpen(buf)

        # 첫 paragraph는 title
        assert doc.paragraphs[0].text == "Report"
        # heading 텍스트가 어딘가에 있어야 함
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Section 1" in all_text
        assert "Body text." in all_text
        assert "Item A" in all_text
        assert "Item B" in all_text
        # 테이블 1개
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "H1"
        assert doc.tables[0].rows[1].cells[1].text == "b"


class TestMakeCreateDocx:
    def test_returns_structured_tool(self):
        tool = make_create_docx(user_id="u-1")
        assert tool.name == "create_docx"
        assert "워드" in tool.description or "문서" in tool.description

    def test_invocation_calls_save(self):
        tool = make_create_docx(user_id="u-1")
        with patch(
            "extension_modules.tools.document.docx_tool.save_to_files"
        ) as mock_save:
            mock_save.return_value = MagicMock(
                file_id="f", filename="r.docx", download_url="/", size_bytes=1
            )
            tool.invoke(
                {
                    "filename": "r",
                    "title": "T",
                    "blocks": [{"type": "paragraph", "text": "hi"}],
                }
            )
            mock_save.assert_called_once()
            assert mock_save.call_args.kwargs["user_id"] == "u-1"
            assert mock_save.call_args.kwargs["filename"] == "r.docx"
            assert mock_save.call_args.kwargs["mime"] == DOCX_MIME


# ── Styling tests ─────────────────────────────────────────────────────────────


class TestDocxStyling:
    def test_table_uses_grid_style_with_bold_header(self):
        content = DocxContent(
            filename="r",
            title="T",
            blocks=[DocxBlock(type="table", table=[["H1", "H2"], ["a", "b"]])],
        )
        buf = _build_docx(content)
        buf.seek(0)
        doc = DocxOpen(buf)
        table = doc.tables[0]
        # Table Grid 스타일 적용 (border 가 자동으로 들어가는 내장 스타일)
        assert table.style.name == "Table Grid"
        # 첫 행은 헤더 → 모든 run 이 bold
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    assert run.bold is True
        # 데이터 행은 bold 없음
        for cell in table.rows[1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    assert run.bold is not True

    def test_table_header_cells_have_background_fill(self):
        # 헤더 셀에 w:shd 요소로 #D9E2F3 배경 적용, 데이터 셀에는 없음
        content = DocxContent(
            filename="r",
            title="T",
            blocks=[DocxBlock(type="table", table=[["H1", "H2"], ["a", "b"]])],
        )
        buf = _build_docx(content)
        buf.seek(0)
        doc = DocxOpen(buf)
        table = doc.tables[0]

        for cell in table.rows[0].cells:
            tc_pr = cell._tc.find(qn("w:tcPr"))
            assert tc_pr is not None
            shd = tc_pr.find(qn("w:shd"))
            assert shd is not None, "헤더 셀에 w:shd 요소 없음"
            assert shd.get(qn("w:fill")) == "D9E2F3"

        for cell in table.rows[1].cells:
            tc_pr = cell._tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                shd = tc_pr.find(qn("w:shd"))
                # 데이터 행: shd 없거나 fill="auto" (기본)
                assert shd is None or shd.get(qn("w:fill")) in (None, "auto")
